import re
import os
import io
import zipfile
import tempfile
import streamlit as st
from pathlib import Path

# ─── Configuración de página ──────────────────────────────────────────────────
st.set_page_config(page_title="CAF – Extractor Dispensas", layout="wide", page_icon="🏦")

# ─── Estilos CAF ──────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;600;700;800&family=Open+Sans:wght@400;600&display=swap');

html, body, [class*="css"] {
    font-family: 'Open Sans', sans-serif;
    background-color: #f4f6f9;
}
.caf-header {
    background: linear-gradient(135deg, #004A8F 0%, #006BB6 100%);
    padding: 28px 36px;
    border-radius: 12px;
    margin-bottom: 28px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    box-shadow: 0 4px 20px rgba(0,74,143,0.25);
}
.caf-header-title {
    color: white;
    font-family: 'Montserrat', sans-serif;
    font-size: 26px;
    font-weight: 800;
    letter-spacing: 0.5px;
    margin: 0;
}
.caf-header-sub {
    color: rgba(255,255,255,0.80);
    font-size: 13px;
    font-weight: 600;
    letter-spacing: 1.5px;
    text-transform: uppercase;
    margin-top: 4px;
}
.section-header {
    background: linear-gradient(90deg, #3A7D44 0%, #4E9D5A 100%);
    color: white;
    font-family: 'Montserrat', sans-serif;
    font-weight: 700;
    font-size: 14px;
    letter-spacing: 1.2px;
    text-transform: uppercase;
    padding: 10px 20px;
    border-radius: 6px;
    margin: 24px 0 0 0;
}
.caf-table {
    width: 100%;
    border-collapse: collapse;
    margin-bottom: 8px;
    border-radius: 8px;
    overflow: hidden;
    box-shadow: 0 2px 12px rgba(0,0,0,0.08);
}
.caf-table tr:nth-child(even) { background-color: #EEF3F9; }
.caf-table tr:nth-child(odd)  { background-color: #ffffff; }
.caf-table td {
    padding: 11px 18px;
    font-size: 13.5px;
    border-bottom: 1px solid #dce6f0;
    vertical-align: top;
}
.caf-table td:first-child {
    background-color: #004A8F;
    color: white;
    font-family: 'Montserrat', sans-serif;
    font-weight: 600;
    font-size: 12.5px;
    width: 240px;
    letter-spacing: 0.3px;
}
.caf-table td:last-child { color: #1a2e45; }
.caf-analysis-table {
    width: 100%;
    border-collapse: collapse;
    border-radius: 8px;
    overflow: hidden;
    box-shadow: 0 2px 12px rgba(0,0,0,0.08);
    margin-bottom: 8px;
}
.caf-analysis-table tr { background-color: #ffffff; border-bottom: 1px solid #dce6f0; }
.caf-analysis-table td {
    padding: 14px 18px;
    font-size: 13.5px;
    color: #1a2e45;
    line-height: 1.7;
    vertical-align: top;
}
.caf-analysis-table td.label {
    background-color: #3A7D44;
    color: white;
    font-family: 'Montserrat', sans-serif;
    font-weight: 700;
    font-size: 12px;
    width: 36px;
    text-align: center;
}
.caf-alert {
    background: #FFF8E1;
    border-left: 5px solid #F5A623;
    border-radius: 6px;
    padding: 14px 18px;
    font-size: 13.5px;
    color: #5a3e00;
    line-height: 1.7;
    box-shadow: 0 2px 8px rgba(245,166,35,0.12);
    margin-bottom: 4px;
}
.img-caption {
    font-size: 12px;
    color: #6b8299;
    font-style: italic;
    margin-top: 4px;
    margin-bottom: 16px;
    font-family: 'Montserrat', sans-serif;
}
div.stButton > button {
    background: linear-gradient(135deg, #004A8F, #006BB6);
    color: white;
    font-family: 'Montserrat', sans-serif;
    font-weight: 700;
    font-size: 14px;
    letter-spacing: 0.8px;
    border: none;
    padding: 12px 36px;
    border-radius: 8px;
    cursor: pointer;
    transition: all 0.2s ease;
    box-shadow: 0 4px 14px rgba(0,74,143,0.3);
    width: 100%;
}
div.stButton > button:hover {
    transform: translateY(-1px);
    box-shadow: 0 6px 18px rgba(0,74,143,0.4);
}
.caf-footer {
    text-align: center;
    color: #8fa3bd;
    font-size: 11.5px;
    margin-top: 40px;
    padding-top: 16px;
    border-top: 1px solid #dce6f0;
    letter-spacing: 0.5px;
}
#MainMenu, footer { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# ─── Header ───────────────────────────────────────────────────────────────────
st.markdown("""
<div class="caf-header">
    <div>
        <div class="caf-header-sub">Gerencia Corporativa de Riesgos · Dirección de Riesgo Soberano</div>
        <div class="caf-header-title">Dispensa &nbsp;|&nbsp; Enmienda</div>
    </div>
    <div style="color:white; font-family:Montserrat; font-size:32px; font-weight:900; letter-spacing:-1px;">CAF</div>
</div>
""", unsafe_allow_html=True)

# ─── Upload ───────────────────────────────────────────────────────────────────
col_up, col_btn = st.columns([3, 1])
with col_up:
    archivo = st.file_uploader("Sube tu documento (.pdf o .docx)", type=["pdf", "docx"])
with col_btn:
    st.write("")
    st.write("")
    procesar = st.button("🔍 Procesar")

# ─── Funciones de extracción de texto ────────────────────────────────────────
def extraer_texto_pdf(ruta):
    import pdfplumber
    texto_completo = []
    with pdfplumber.open(ruta) as pdf:
        for pagina in pdf.pages:
            for tabla in pagina.extract_tables():
                for fila in tabla:
                    texto_completo.append(" | ".join([c or "" for c in fila]))
            texto_plano = pagina.extract_text()
            if texto_plano:
                texto_completo.append(texto_plano)
    return "\n".join(texto_completo)

def extraer_texto_docx(ruta):
    from docx import Document
    doc = Document(ruta)
    lineas = []
    for parrafo in doc.paragraphs:
        if parrafo.text.strip():
            lineas.append(parrafo.text.strip())
    for tabla in doc.tables:
        for fila in tabla.rows:
            lineas.append(" | ".join([c.text.strip() for c in fila.cells]))
    return "\n".join(lineas)

# ─── Extracción de imágenes ───────────────────────────────────────────────────
def extraer_imagenes_pdf(ruta, palabras_clave):
    """Devuelve imágenes de páginas que contengan las palabras clave."""
    imagenes = []
    try:
        import fitz
        doc = fitz.open(ruta)
        for num_pagina, pagina in enumerate(doc):
            texto_pagina = pagina.get_text().lower()
            if any(p.lower() in texto_pagina for p in palabras_clave):
                mat = fitz.Matrix(2, 2)
                pix = pagina.get_pixmap(matrix=mat)
                imagenes.append({
                    "label": f"Página {num_pagina + 1}",
                    "bytes": pix.tobytes("png"),
                })
        doc.close()
    except ImportError:
        st.info("ℹ️ Instala `PyMuPDF` en requirements.txt para ver imágenes de PDF.")
    return imagenes

def extraer_imagenes_docx(ruta):
    """Extrae todas las imágenes embebidas en el .docx."""
    imagenes = []
    try:
        with zipfile.ZipFile(ruta, "r") as z:
            nombres = sorted([n for n in z.namelist() if n.startswith("word/media/")])
            for nombre in nombres:
                ext_img = Path(nombre).suffix.lower()
                if ext_img in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"]:
                    imagenes.append({
                        "label": nombre.split("/")[-1],
                        "bytes": z.read(nombre),
                    })
    except Exception as e:
        st.warning(f"Error extrayendo imágenes del docx: {e}")
    return imagenes

def mostrar_imagenes(imagenes, titulo):
    if not imagenes:
        return
    st.markdown(f'<div class="section-header">🖼️ &nbsp;{titulo}</div>', unsafe_allow_html=True)
    st.write("")
    for img in imagenes:
        col_izq, col_centro, col_der = st.columns([1, 1, 1])
        with col_centro:
            st.image(img["bytes"], use_column_width=True)
            st.markdown(f'<div class="img-caption" style="text-align:center">{img["label"]}</div>', unsafe_allow_html=True)

# ─── Funciones de extracción de campos 
def extraer_dispensa_si(ruta, extension):
    filas_tabla = []
    if extension == ".pdf":
        import pdfplumber
        with pdfplumber.open(ruta) as pdf:
            for pagina in pdf.pages:
                for tabla in pagina.extract_tables():
                    tabla_flat = " ".join([str(c) for fila in tabla for c in fila if c])
                    if "Dispensa" in tabla_flat and "Instancia" in tabla_flat:
                        filas_tabla = tabla
                        break
    elif extension == ".docx":
        from docx import Document
        doc = Document(ruta)
        for tabla in doc.tables:
            tabla_flat = " ".join([c.text for fila in tabla.rows for c in fila.cells])
            if "Dispensa" in tabla_flat and "Instancia" in tabla_flat:
                filas_tabla = [[c.text.strip() for c in fila.cells] for fila in tabla.rows]
                break
    if not filas_tabla:
        return []
    resultados_si = []
    for fila in filas_tabla[1:]:
        if len(fila) < 4:
            continue
        tipo         = fila[0].strip() if fila[0] else ""
        instancia    = fila[2].strip() if fila[2] else ""
        subtipo      = fila[1].strip() if fila[1] else ""
        marca_si     = fila[3].strip() if fila[3] else ""
        es_extension = "Extensi" in tipo
        if any(x in marca_si for x in ["x", "X", "✗", "✘", "☒"]):
            resultados_si.append({
                "Tipo"     : subtipo if es_extension else tipo,
                "Instancia": instancia,
            })
    return resultados_si

def extraer_extension_plazo(texto):
    match = re.search(r"((?:X|x|✗|✘|☒)\s*Hasta\s+\d+\s+meses\s*\([^)]+\))", texto, re.IGNORECASE)
    if match:
        return re.sub(r"^[Xx✗✘☒]\s*", "", match.group(1).strip())
    return "No encontrado"

def buscar_campo(texto, patron, grupo=1):
    match = re.search(patron, texto, re.IGNORECASE | re.MULTILINE)
    return match.group(grupo).strip().strip("|").strip() if match else "No encontrado"

def extraer_parrafo_pendiente(texto):
    encontrados = []
    for parrafo in re.split(r"\n{2,}", texto):
        if "pendiente por justificar" in parrafo.lower():
            limpio = " ".join(parrafo.split()).strip()
            if limpio:
                encontrados.append(limpio)
    return encontrados if encontrados else None

def extraer_seccion_justificacion(texto):
    lineas = texto.split("\n")
    inicio = fin = None
    for i, linea in enumerate(lineas):
        if re.search(r"^Justificaci[oó]n\s*$", linea.strip(), re.IGNORECASE):
            inicio = i + 1
        if inicio and re.search(r"^Recomendaci[oó]n\s*$", linea.strip(), re.IGNORECASE):
            fin = i
            break
    if inicio is None:
        return None
    contenido = "\n".join(lineas[inicio:fin]).strip()
    return re.sub(r"\n{3,}", "\n\n", contenido)

def tabla_html(filas):
    rows = ""
    for label, valor in filas:
        rows += f"<tr><td>{label}</td><td>{valor or '—'}</td></tr>"
    return f'<table class="caf-table">{rows}</table>'

# ─── Procesamiento ────────────────────────────────────────────────────────────
if procesar:
    if archivo is None:
        st.warning("⚠️ Por favor sube un archivo antes de procesar.")
    else:
        extension = Path(archivo.name).suffix.lower()

        with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as tmp:
            tmp.write(archivo.read())
            ruta_tmp = tmp.name

        with st.spinner("Procesando documento..."):
            texto     = extraer_texto_pdf(ruta_tmp) if extension == ".pdf" else extraer_texto_docx(ruta_tmp)
            dispensas = extraer_dispensa_si(ruta_tmp, extension)
            pendiente = extraer_parrafo_pendiente(texto)
            justif    = extraer_seccion_justificacion(texto)

            # Imágenes
            if extension == ".pdf":
                imgs_pendiente = extraer_imagenes_pdf(ruta_tmp, [
                    "justificativos pendientes", "cuadro n°2", "cuadro no 2",
                    "pendiente por justificar", "justificativo"
                ])
                imgs_justif = extraer_imagenes_pdf(ruta_tmp, [
                    "proyección de utilizaciones", "cuadro n°3", "cuadro no 3",
                    "pagos recursos pendientes", "proyeccion de utilizaciones"
                ])
            else:
                todas_imgs     = extraer_imagenes_docx(ruta_tmp)
                mid            = len(todas_imgs) // 2
                imgs_pendiente = todas_imgs[:mid]      # primera mitad → monto por justificar
                imgs_justif    = todas_imgs[mid:mid+1] # solo 1 imagen → justificación

        os.unlink(ruta_tmp)

        nombre_prestatario = buscar_campo(texto, r"Nombre del Prestatario[\s|]*([^\n|]+)")
        nombre_operacion   = buscar_campo(texto, r"Nombre de la Operaci[oó]n[\s|]*([^\n|]+)")
        monto_aprobado     = buscar_campo(texto, r"(?<![a-zA-Z])Aprobado[\s|]*([\d.,]+(?:\s*MM)?)")
        monto_desembolsado = buscar_campo(texto, r"(?<![a-zA-Z])Desembolsado[\s|]*([\d.,]+(?:\s*(?:\([^)]+\))?\s*MM)?)")
        unidad_negocios    = buscar_campo(texto, r"Gerencia de Negocios[\s|]*([^\n|]+)")
        garante            = buscar_campo(texto, r"Garante[\s|]*([^\n|]+)")
        extension_plazo    = extraer_extension_plazo(texto)
        tipo_dispensa      = dispensas[0]["Tipo"]      if dispensas else "Ninguna marcada con SI"
        instancia_aprob    = dispensas[0]["Instancia"] if dispensas else "—"

        # ── Tabla 1: Informe de la Operación ──────────────────────────────────
        st.markdown('<div class="section-header">📋 &nbsp;Informe de la Operación</div>', unsafe_allow_html=True)
        st.markdown(tabla_html([
            ("Nombre del Prestatario",    nombre_prestatario),
            ("Nombre de la Operación",    nombre_operacion),
            ("Monto Aprobado",            monto_aprobado),
            ("Monto Desembolsado",        monto_desembolsado),
            ("Unidad de Negocios",        unidad_negocios),
            ("Garante",                   garante),
            ("Instancia de Aprobación",   extension_plazo),
            ("Tipo de Dispensa/Enmienda", tipo_dispensa),
            ("Instancia Aprobatoria",     instancia_aprob),
        ]), unsafe_allow_html=True)

        # ── Tabla 2: Monto por Justificar ─────────────────────────────────────
        st.markdown('<div class="section-header">⚠️ &nbsp;Monto por Justificar</div>', unsafe_allow_html=True)
        if pendiente:
            rows = "".join([f"<tr><td>{p}</td></tr>" for p in pendiente])
            st.markdown(f'<table class="caf-analysis-table">{rows}</table>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="caf-alert">No se encontró la frase <strong>\'pendiente por justificar\'</strong>.</div>', unsafe_allow_html=True)

        mostrar_imagenes(imgs_pendiente, "Cuadros – Monto por Justificar")

        # ── Tabla 3: Justificación ─────────────────────────────────────────────
        st.markdown('<div class="section-header">📝 &nbsp;Justificación</div>', unsafe_allow_html=True)
        if justif:
            parrafos = [p.strip() for p in justif.split("\n") if p.strip()]
            rows = ""
            contador = 1
            for p in parrafos:
                if re.match(r"^(i{1,3}v?|vi{0,3}|ix|x)\)", p, re.IGNORECASE):
                    rows += f'<tr><td class="label">{contador}</td><td>{p}</td></tr>'
                    contador += 1
                else:
                    rows += f'<tr><td colspan="2">{p}</td></tr>'
            st.markdown(f'<table class="caf-analysis-table">{rows}</table>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="caf-alert">No se encontró la sección de <strong>Justificación</strong>.</div>', unsafe_allow_html=True)

        mostrar_imagenes(imgs_justif, "Cuadros – Justificación")

        st.markdown('<div class="caf-footer">CAF – Banco de Desarrollo de América Latina y el Caribe &nbsp;·&nbsp; Gerencia Corporativa de Riesgos</div>', unsafe_allow_html=True)




